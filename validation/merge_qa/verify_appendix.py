import re
from docx import Document

document = Document("validation/merge_qa/merged-9000-v2.docx")
paragraphs = [paragraph.text for paragraph in document.paragraphs]
ignore_index = paragraphs.index("IGNORE")
pattern = re.compile(r"[\w]+(?:['’.-][\w]+)*", re.UNICODE)
appendix_text = " ".join(paragraphs[ignore_index + 1:])
all_text = " ".join(paragraphs)
print("appendix_words", len(pattern.findall(appendix_text)))
print("ignore", paragraphs.count("IGNORE"))
print("removed_references", all_text.count("Removed Author"), all_text.count("Delete Me"))
print("tables", len(document.tables), "images", len(document.inline_shapes), "paragraphs", len(paragraphs))
